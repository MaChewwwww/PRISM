from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OBSIDIAN = "0B0F14"
TEAL = "547D83"
SLATE = "475569"
PALE = "EEF4F4"
LIGHT_GRAY = "F2F4F7"
FONT_SANS = "Calibri"
FONT_SERIF = "Georgia"
FONT_MONO = "Cascadia Mono"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _shade(cell: object, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_margins(cell: object) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), "80" if edge in {"top", "bottom"} else "120")
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row: object) -> None:
    properties = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _page_number(paragraph: object) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _configure_list_numbering(document: DocumentType) -> None:
    numbering = document.part.numbering_part.element
    for style_name, marker in (("List Bullet", "\u2022"), ("List Number", "%1.")):
        style = document.styles[style_name]
        number_id = style._element.find(".//" + qn("w:numId"))
        if number_id is None:
            raise ValueError(f"{style_name} must use a real numbering definition")
        number_value = number_id.get(qn("w:val"))
        number = next(
            node
            for node in numbering.findall(qn("w:num"))
            if node.get(qn("w:numId")) == number_value
        )
        abstract_value = number.find(qn("w:abstractNumId")).get(qn("w:val"))
        abstract = next(
            node
            for node in numbering.findall(qn("w:abstractNum"))
            if node.get(qn("w:abstractNumId")) == abstract_value
        )
        level = next(
            node
            for node in abstract.findall(qn("w:lvl"))
            if node.get(qn("w:ilvl")) == "0"
        )
        level.find(qn("w:lvlText")).set(qn("w:val"), marker)

        properties = level.find(qn("w:pPr"))
        if properties is None:
            properties = OxmlElement("w:pPr")
            level.append(properties)
        indent = properties.find(qn("w:ind"))
        if indent is None:
            indent = OxmlElement("w:ind")
            properties.append(indent)
        indent.set(qn("w:left"), "720")
        indent.set(qn("w:hanging"), "360")

        tabs = properties.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            properties.append(tabs)
        for child in list(tabs):
            tabs.remove(child)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)

        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            properties.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")

        run_properties = level.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            level.append(run_properties)
        fonts = run_properties.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.append(fonts)
        fonts.set(qn("w:ascii"), FONT_SANS)
        fonts.set(qn("w:hAnsi"), FONT_SANS)


def _font(run: object, name: str, size: float, color: str, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^]]+\]\([^)]+\))")


def _hyperlink(paragraph: object, label: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((run_properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _inline(paragraph: object, text: str) -> None:
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("["):
            link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", token)
            if link:
                _hyperlink(paragraph, link.group(1), link.group(2))
        else:
            run = paragraph.add_run(token[1:-1])
            _font(run, FONT_MONO, 9, OBSIDIAN)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E7EEEE")
            run._r.get_or_add_rPr().append(shading)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _configure(document: DocumentType, source: Path) -> None:
    # standard_business_brief preset with named PRISM color/title overrides.
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT_SANS
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_SANS)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_SANS)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(OBSIDIAN)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for name, size, color, before, after in (
        ("Title", 30, OBSIDIAN, 0, 6),
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, SLATE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = FONT_SANS if name != "Title" else FONT_SERIF
        family = FONT_SANS if name != "Title" else FONT_SERIF
        style._element.rPr.rFonts.set(qn("w:ascii"), family)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), family)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), family)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = FONT_SANS
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_SANS)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_SANS)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    _configure_list_numbering(document)

    if "PRISM Code" not in [style.name for style in document.styles]:
        code = document.styles.add_style("PRISM Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = document.styles["PRISM Code"]
    code.font.name = FONT_MONO
    code._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)
    code.font.size = Pt(8.6)
    code.font.color.rgb = RGBColor.from_string(OBSIDIAN)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.1

    if "PRISM Contents" not in [style.name for style in document.styles]:
        contents = document.styles.add_style("PRISM Contents", WD_STYLE_TYPE.PARAGRAPH)
    else:
        contents = document.styles["PRISM Contents"]
    contents.font.name = FONT_SANS
    contents._element.rPr.rFonts.set(qn("w:ascii"), FONT_SANS)
    contents._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_SANS)
    contents.font.size = Pt(9.5)
    contents.font.color.rgb = RGBColor.from_string(OBSIDIAN)
    contents.paragraph_format.left_indent = Inches(0.16)
    contents.paragraph_format.space_after = Pt(2)
    contents.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PRISM  /  PROJECT CONCEPT")
    _font(run, FONT_SANS, 8, TEAL, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ecosystem-consolidation-v1  /  ")
    _font(run, FONT_SANS, 8, SLATE)
    _page_number(footer)

    document.core_properties.title = "PRISM Project Concept"
    document.core_properties.subject = (
        "PM-oriented product concept for governed paper trading"
    )
    document.core_properties.author = "PRISM"
    document.core_properties.comments = f"Generated from {source.as_posix()}"


def _cover(document: DocumentType, revision: str) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("PROJECT CONCEPT")
    _font(run, FONT_SANS, 9, TEAL, bold=True)

    title = document.add_paragraph("PRISM", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("A governed, paper-only market intelligence platform")
    _font(run, FONT_SANS, 14, SLATE, bold=True)

    tagline = document.add_paragraph()
    tagline.paragraph_format.space_after = Pt(18)
    run = tagline.add_run("One signal. Multiple perspectives. Better decisions.")
    _font(run, FONT_SANS, 10.5, TEAL, bold=True)

    note = document.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.16)
    note.paragraph_format.right_indent = Inches(0.16)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(14)
    note_properties = note._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE)
    note_properties.append(shading)
    _inline(
        note,
        "AI broadens the analysis. Deterministic rules control authorization. "
        "Every order is Alpaca paper-only; live trading is prohibited.",
    )

    for label, value in (
        ("Audience", "Project manager, reviewers, and delivery team"),
        ("Scope", "Complete product concept; delivery progress is tracked separately"),
        ("Revision", revision),
    ):
        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        label_run = meta.add_run(f"{label}: ")
        _font(label_run, FONT_SANS, 9.5, OBSIDIAN, bold=True)
        value_run = meta.add_run(value)
        _font(value_run, FONT_SANS, 9.5, SLATE)


def _contents(document: DocumentType, headings: list[str]) -> None:
    document.add_heading("Contents", level=1)
    for heading in headings:
        paragraph = document.add_paragraph(style="PRISM Contents")
        text_run = paragraph.add_run(heading)
        _font(text_run, FONT_SANS, 9.5, OBSIDIAN)
    document.add_page_break()


def _column_widths(rows: list[list[str]]) -> list[int]:
    header = rows[0]
    count = len(header)
    if count == 2:
        first = header[0].strip()
        if first == "Surface":
            return [2000, 7360]
        if first in {"Area", "Control"}:
            return [2200, 7160]
        return [2600, 6760]
    if count == 3:
        if header[0].strip() == "Audience":
            return [1500, 3300, 4560]
        return [2300, 5100, 1960]
    if count == 5:
        return [1600, 1900, 2200, 1800, 1860]
    base, remainder = divmod(CONTENT_WIDTH_DXA, count)
    return [base + (1 if index < remainder else 0) for index in range(count)]


def _set_table_geometry(table: object, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr  # type: ignore[attr-defined]

    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid  # type: ignore[attr-defined]
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def _table(document: DocumentType, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cell = cells[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _inline(paragraph, value.strip())
            for run in paragraph.runs:
                _font(
                    run,
                    FONT_SANS,
                    9.2,
                    TEAL if row_index == 0 else OBSIDIAN,
                    bold=row_index == 0 or column == 0,
                )
            if row_index == 0:
                _shade(cell, LIGHT_GRAY)
        if row_index == 0:
            _repeat_table_header(table.rows[-1])
    _set_table_geometry(table, _column_widths(rows))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _parse(markdown: str, document: DocumentType) -> None:
    lines = markdown.splitlines()
    revision_match = re.search(r"Revision: `([^`]+)`", markdown)
    revision = revision_match.group(1) if revision_match else "unversioned"
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    _cover(document, revision)
    _contents(document, headings)

    paragraph_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        _inline(paragraph, " ".join(paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                paragraph = document.add_paragraph(style="PRISM Code")
                paragraph.paragraph_format.keep_together = True
                run = paragraph.add_run("\n".join(code_buffer))
                _font(run, FONT_MONO, 8.6, OBSIDIAN)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "EEF4F4")
                paragraph._p.get_or_add_pPr().append(shading)
                code_buffer.clear()
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_buffer.append(line)
            index += 1
            continue
        if (
            line.startswith("| ")
            and index + 1 < len(lines)
            and lines[index + 1].startswith("| ---")
        ):
            flush_paragraph()
            rows: list[list[str]] = []
            rows.append([part.strip() for part in line.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(
                    [part.strip() for part in lines[index].strip("|").split("|")]
                )
                index += 1
            _table(document, rows)
            continue
        if not line:
            flush_paragraph()
            index += 1
            continue
        if line.startswith(("**One signal.", "Revision: `")):
            flush_paragraph()
            index += 1
            continue
        if line.startswith(("# ", "Revision:")):
            flush_paragraph()
        elif line.startswith("## "):
            flush_paragraph()
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            flush_paragraph()
            document.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\. ", line):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            _inline(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            _inline(paragraph, line[2:])
        else:
            paragraph_buffer.append(line)
        index += 1
    flush_paragraph()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    markdown = args.source.read_text(encoding="utf-8")
    document = Document()
    _configure(document, args.source)
    _parse(markdown, document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
